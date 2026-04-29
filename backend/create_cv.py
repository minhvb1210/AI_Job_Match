from docx import Document
import os

def create_cv():
    doc = Document()
    doc.add_heading('Demo Candidate CV', 0)
    doc.add_paragraph('Experienced Backend Developer')
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, FastAPI, Docker, PostgreSQL, Flutter, AI, Machine Learning')
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Developed AI recruitment platforms and optimized backend services.')
    doc.save('cv.docx')
    print("cv.docx created successfully.")

if __name__ == "__main__":
    create_cv()
